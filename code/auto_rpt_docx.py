# -*- coding: utf-8 -*-
"""
Created on Mon Aug 29 23:43:17 2022

@author: lxl
"""

from auto_rpt_func import *
from docx import Document
from docx.shared import Inches, Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING

from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

import pandas as pd
import os

# 全文字体设置
def docxinitial(document):
    # 设置正文字体类型、大小
    document.styles["Normal"].font.name = u'Times New Roman'
    document.styles["Normal"].font.size = Pt(16) # 
    document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

# 大标题设置
def title0(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36) 
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(18)
    ti1.font.name = u'黑体'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 时间行设置
def time_note(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    ti2 = ti.add_run('\n')
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 一级标题设置
def title1(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 二级标题设置
def title2(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.bold = True
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 单一段落格式设置
def one_paragraph_txt(document, txt):
    txt1 = txt.split("，")[0]
    txt2 = txt.replace(txt1,'')
    p = document.add_paragraph()
    #p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(36)
    p.style.font.size = Pt(16)
    p1 = p.add_run(txt1)
    p1.font.size = Pt(16)
    p1.font.name = u'Times New Roman'
    r = p1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p2 = p.add_run(txt2)
    p2.font.size = Pt(16)
    p.paragraph_format.first_line_indent = p.style.font.size * 2 # 缩进2字符
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    
# 多段格式设置
def much_paragraph_txt(document, txt):
    txt_list = txt.split("\n")
    for t in txt_list:
        one_paragraph_txt(document, t)

# 页码设置
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10.5)

# 生成文件
def create_file(date_str,data,config_data,data_dim,title0_name='产业上中下游价格指数周报',file_pre='产业上中下游价格指数周报',file_dir='1'):
    time_str = datetime.datetime.strptime(date_str,'%Y-%m-%d')
    begin_date = (time_str-datetime.timedelta(days=time_str.weekday())).strftime("%Y年%m月%d日")
    end_date = (time_str+datetime.timedelta(days=6-time_str.weekday())).strftime("%Y年%m月%d日")
    
    document = Document()
    docxinitial(document)
    
    # title0(document, '产业上中下游价格指数周报')
    title0(document, title0_name)
    time_note(document, "（"+begin_date+"-"+end_date+"）")
    
    df = config_data[['first_level_para_id','second_level_para_id']].drop_duplicates()
    df_list = df.apply(lambda x:tuple(x),axis = 1).values.tolist() 
    first_ti = list(set(df.loc[:,'first_level_para_id'].tolist()))
    second_ti = list(set(df.loc[:,'second_level_para_id'].tolist()))
    
    # 一级标题
    for x in range(0,len(first_ti)):
        if first_ti[x] == 1:
            text1 = '一、'
        if first_ti[x] == 2:
            text1 = '二、'
        if first_ti[x] == 3:
            text1 = '三、'
        if first_ti[x] == 4:
            text1 = '四、'
        if first_ti[x] == 5:
            text1 = '五、'   
        if first_ti[x] == 6:
            text1 = '六、' 
        if first_ti[x] == 7:
            text1 = '七、' 
        title1(document, text1 + config_data[config_data['first_level_para_id'] == first_ti[x]]['first_level_para_name'].tolist()[0])
        # 二级标题
        for y in range(0,len(second_ti)):
            if second_ti[y] == 1:
                text2 = '（一）'
            if second_ti[y] == 2:
                text2 = '（二）'
            if second_ti[y] == 3:
                text2 = '（三）'
            if second_ti[y] == 4:
                text2 = '（四）'
            if second_ti[y] == 5:
                text2 = '（五）'   
            if second_ti[y] == 6:
                text2 = '（六）'
            if second_ti[y] == 7:
                text2 = '（七）'
            if (first_ti[x],second_ti[y]) in df_list:
                title2(document, text2 + config_data[(config_data['first_level_para_id'] == first_ti[x]) & (config_data['second_level_para_id'] == second_ti[y])]['second_level_para_name'].tolist()[0])
                much_paragraph_txt(document, get_paragraph(config_data,data_dim,date_str,data,first_ti[x],second_ti[y]))
    
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] # 获取页眉的第一个段落 
    #paragraph.add_run('这是第一节的页眉') # 添加页面内容
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(paragraph.add_run())
       # 构建文件名和路径
    file_name = f'{file_pre}（{begin_date}-{end_date}）.docx'
    file_path = f'/Users/harvin/code/自动报告产品开发-产业链@20220830/output/{file_dir}/{file_name}'
    # 创建目录（如果不存在）
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # 保存文件
    document.save(file_path)
    print(file_path)


