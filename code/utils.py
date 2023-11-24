from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime
from io import BytesIO

class WordReportGenerator:
    def __init__(self):
        self.document = Document()
        self.docxinitial()

    def docxinitial(self):
        self.document.styles["Normal"].font.name = 'Times New Roman'
        self.document.styles["Normal"].font.size = Pt(16)
        self.document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    
    # 大标题设置
    def title0(self, txt):
        title = txt
        ti = self.document.add_paragraph()
        ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ti.paragraph_format.space_before = Pt(0)
        ti.paragraph_format.space_after = Pt(0)
        ti.paragraph_format.line_spacing = Pt(36) 
        ti1 = ti.add_run(title)
        ti1.font.size = Pt(18)
        ti1.font.name = '黑体'
        r = ti1._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    #题注：宋体小四居中1.5倍行距
    def little_title(self, txt):
        if txt:
            title = txt
            ti = self.document.add_paragraph()
            ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ti.paragraph_format.space_before = Pt(0)
            ti.paragraph_format.space_after = Pt(0)
            ti.paragraph_format.line_spacing = Pt(36)
            ti1 = ti.add_run(title)
            ti1.font.size = Pt(12)  # 设置字体大小为小四
            ti1.font.name = '宋体'
            r = ti1._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距

    # 时间行设置
    def time_note(self, txt):
        if txt:
            title = txt
            ti = self.document.add_paragraph()
            ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ti.paragraph_format.space_before = Pt(0)
            ti.paragraph_format.space_after = Pt(0)
            ti.paragraph_format.line_spacing = Pt(36)
            ti1 = ti.add_run(title)
            ti1.font.size = Pt(16)
            ti1.font.name = u'Times New Roman'
            r = ti1._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            ti2 = ti.add_run('\n')
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

    # 一级标题设置
    def title1(self, txt):
        if txt:
            title = txt
            ti = self.document.add_paragraph()
            ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ti.paragraph_format.space_before = Pt(0)
            ti.paragraph_format.space_after = Pt(0)
            ti.paragraph_format.line_spacing = Pt(36)
            ti1 = ti.add_run(title)
            ti1.font.size = Pt(16)
            ti1.font.name = u'Times New Roman'
            r = ti1._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    # 二级标题设置
    def title2(self, txt):
        if txt:
            title = txt
            ti = self.document.add_paragraph()
            ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ti.paragraph_format.space_before = Pt(0)
            ti.paragraph_format.space_after = Pt(0)
            ti.paragraph_format.line_spacing = Pt(36)
            ti1 = ti.add_run(title)
            ti1.font.bold = True
            ti1.font.size = Pt(16)
            ti1.font.name = u'Times New Roman'
            r = ti1._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
            ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

    # 单一段落格式设置 
    def one_paragraph_txt(self, txt):
        if txt:
            txt1 = txt.split("，")[0]
            txt2 = txt.replace(txt1,'')
            p = self.document.add_paragraph()
            #p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(36)
            p.style.font.size = Pt(16)
            p1 = p.add_run(txt1)
            p1.font.size = Pt(16)
            p1.font.name = u'Times New Roman'
            r = p1._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p2 = p.add_run(txt2)
            p2.font.size = Pt(16)
            p.paragraph_format.first_line_indent = p.style.font.size * 2 # 缩进2字符
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    # 格式化段落（无特殊格式，只有大小和罗马字符）JUSTIFY
    def standard_paragraph_txt(self, txt):
        if txt:
            p = self.document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(36)
            p.style.font.size = Pt(16)
            
            run = p.add_run(txt)
            run.font.size = Pt(16)
            for char in txt:
                if char.isdigit() or char.isalpha():  # 检查字符是否为数字或英文字母
                    run.font.name = 'Times New Roman'  # 设置数字和英文字符的字体为 Times New Roman
            
            p.paragraph_format.first_line_indent = p.style.font.size * 2
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 多段格式设置
    def much_paragraph_txt(self, txt):
        txt_list = txt.split("\n")
        for txt in txt_list:
            self.one_paragraph_txt(txt)

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(qn(name), value)

    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(10.5)

    def add_image_to_docx(self, image_path, width=Inches(6)):
        self.document.add_picture(image_path, width=width)  # 添加指定宽度的图片
        # 设置最后一个段落居中对齐
        paragraphs = self.document.paragraphs
        if paragraphs:
            paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_table_from_dataframe(self, dataframe, font_size=12, align_center=True):
        table = self.document.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'
        table.autofit = True  # 自适应内容宽度
        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(dataframe.columns):
            hdr_cells[i].text = str(column_name)
            # 设置表头的字体样式和居中
            cell = hdr_cells[i]
            cell.paragraphs[0].alignment = 1  # 1代表居中对齐
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(font_size)
                    if cell == hdr_cells[i]:  # 判断是否为表头单元格
                        run.font.bold = True  # 设置表头字体加粗

        # 添加数据行
        for index, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = row_cells[i]
                cell.text = str(value)
                # 设置每个单元格的字体样式和居中
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(font_size)
                        paragraph.alignment = 1 if align_center else 0  # 设置居中或左对齐



    def create_file(self, date_str, data, data_dim, title0_name='信息软件业招聘岗位数分析报告',time_note='202307', file_pre='信息软件业招聘岗位数分析报告', file_dir='1'):
        time_str = datetime.datetime.strptime(date_str, '%Y-%m-%d')
        begin_date = (time_str - datetime.timedelta(days=time_str.weekday())).strftime("%Y年%m月%d日")
        end_date = (time_str + datetime.timedelta(days=6 - time_str.weekday())).strftime("%Y年%m月%d日")

        self.title0(title0_name)
        # self.time_note( "（" + begin_date + "-" + end_date + "）")
        self.time_note(time_note)

        # 这里用 data 来模拟 config_data
        # 假设 data 是一段字符串列表，每个字符串代表一个标题或段落
        # 可以按需修改此部分逻辑以适应实际的字符串处理

        titles = [title.strip() for title in data.split('\n') if title.strip()]
        for title in titles:
            self.title1(title)

        footer = self.document.sections[0].footer
        paragraph = footer.paragraphs[0]  # 获取页眉的第一个段落
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_page_number(paragraph.add_run())

        file_name = f'{file_pre}（{begin_date}-{end_date}）.docx'
        file_name=f'{file_pre}{time_note}.docx'
        file_path = f'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/output/{file_dir}/{file_name}'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        self.document.save(file_path)
        print(file_path)


    def job_docx_file(self,title0,title1,pic_title,table_title,par1,pars,file_pre,file_dir,pic_dir,df_table,year_month):
        self.title0(title0)
        self.title1(title1)
        self.time_note(year_month)
        self.standard_paragraph_txt(par1)
        self.add_image_to_docx(pic_dir)
        self.little_title(pic_title)
        self.standard_paragraph_txt(pars)
        self.little_title(table_title)
        self.add_table_from_dataframe(df_table)
        file_name=f'{file_pre}{year_month}.docx'
        file_path = f'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/output/{file_dir}/{file_pre}/{file_name}'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        self.document.save(file_path)
        print(file_path)


    def job_salary_docx_file(self,title0,pic_title,par1,file_pre,file_dir,pic_dir,year_month):
        self.title0(title0)
        self.time_note(year_month)
        self.standard_paragraph_txt(par1)
        self.add_image_to_docx(pic_dir)
        self.little_title(pic_title)

        file_name=f'{file_pre}{year_month}.docx'
        file_path = f'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/output/{file_dir}/{file_pre}/{file_name}'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        self.document.save(file_path)
        print(file_path)

    def ie_fina_docx_file(self,title0,par1,par2,par3,pic_dir,pic_dir2,pic_dir3,file_pre,file_dir,year_month,lt,lt2,lt3):
        self.title0(title0)
        self.time_note(year_month)
        self.standard_paragraph_txt(par1)
        
        self.standard_paragraph_txt(par2)
        self.add_image_to_docx(pic_dir2)
        self.little_title(lt2)
        
        self.standard_paragraph_txt(par3)
        self.add_image_to_docx(pic_dir)
        self.little_title(lt)
        self.add_image_to_docx(pic_dir3)
        self.little_title(lt3)
        

        file_name=f'{file_pre}{year_month}.docx'
        file_path = f'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/output/{file_dir}/{file_pre}/{file_name}'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        self.document.save(file_path)
        print(file_path)

def longest_common_substring(s1, s2):
    # 基于最大重复子串匹配
    m = [[0] * (1 + len(s2)) for _ in range(1 + len(s1))]
    longest, x_longest = 0, 0
    for x in range(1, 1 + len(s1)):
        for y in range(1, 1 + len(s2)):
            if s1[x - 1] == s2[y - 1]:
                m[x][y] = m[x - 1][y - 1] + 1
                if m[x][y] > longest:
                    longest = m[x][y]
                    x_longest = x
            else:
                m[x][y] = 0
    return s1[x_longest - longest: x_longest]

def get_best_match(name, ordered_names):
    max_common_substring = 0
    best_match = None
    for ordered_name in ordered_names:
        common_substring = longest_common_substring(name, ordered_name)
        if len(common_substring) > max_common_substring:
            max_common_substring = len(common_substring)
            best_match = ordered_name
    return best_match