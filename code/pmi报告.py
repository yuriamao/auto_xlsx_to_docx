# -*- coding: utf-8 -*-
"""
Created on Wed Nov 22 13:44:23 2023

@author: lxl
"""
import sys
import os
import numpy as np
import pandas as pd
import locale
import datetime
import math

import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from matplotlib.dates import DateFormatter, MonthLocator  # Add this line
from matplotlib.ticker import FuncFormatter # 用于改y轴格式
from docx import Document
from io import BytesIO

import matplotlib

from docx.enum.table import WD_TABLE_ALIGNMENT

#cu = os.path.abspath(os.path.join(os.getcwd(), ".."))
#os.chdir(cu)
locale.setlocale(locale.LC_CTYPE, 'chinese')

# 数据处理 
filename = r'.\data\指标数据20231120.xlsx'
data_df = pd.read_excel(r'.\data\指标数据20231120.xlsx',header=0,skiprows = range(1,2))
pmi = data_df[(data_df['指标名称'] == 'PMI指数_当期值') & (data_df['产业名称'] == '制造业采购经理指数')]
pmi_other = data_df[(data_df['指标名称'] == 'PMI指数_当期值')]


# 与全国对比 
def calc_country_change(df,period):
    country = float(df[((df['地区编码'] == 0) | (df['地区编码'] == '000000')) & (df['报告期编码'] == period)]['数据值'])
    beijing = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == period)]['数据值'])
    diff = round(beijing-country,1)
    if diff > 0:
        judge1 = '与同期全国制造业PMI相比高'+str(abs(diff))+'个百分点'
    if diff < 0:
        judge1 = '与同期全国制造业PMI相比低'+str(abs(diff))+'个百分点'
    if diff == 0:
        judge1 = '与同期全国制造业PMI相比持平'
    return judge1

# 与全国对比2
def calc_country_change2(df,period):
    country = float(df[((df['地区编码'] == 0) | (df['地区编码'] == '000000')) & (df['报告期编码'] == period)]['数据值'])
    beijing = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == period)]['数据值'])
    diff = round(beijing-country,1)
    if diff > 0:
        judge4 = '较全国高'+str(abs(diff))+'个百分点'
    if diff < 0:
        judge4 = '较全国低'+str(abs(diff))+'个百分点'
    if diff == 0:
        judge4 = '与全国持平'
    return judge4

# 计算同比
def calc_tongbi(df,period):
    this_month = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == period)]['数据值'])
    last_period = str(int(period[0:4])-1) + period[5:7].zfill(2)
    this_month_last_year = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == last_period)]['数据值'])
    tongbi = round((this_month-this_month_last_year)/this_month_last_year * 100,1)
    if tongbi > 0:
        judge2 = '同比上升'+str(abs(tongbi))+'个百分点'
    if tongbi < 0:
        judge2 = '同比下降'+str(abs(tongbi))+'个百分点'
    if tongbi == 0:
        judge2 = '同比持平'
    return judge2 

# 计算环比 
def calc_huanbi(df,period):
    this_month = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == period)]['数据值'])
    if period[5:6] == '01':
        last_period = str(period[0:4]-1) + '12'
    else:
        last_period = period[0:4] + str(int(period[5:7])-1).zfill(2)

    last_month = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == last_period)]['数据值'])
    huanbi = round((this_month-last_month)/last_month * 100,1)
    if huanbi > 0:
        judge3 = '较上月上升'+str(abs(huanbi))+'个百分点'
    if huanbi < 0:
        judge3 = '较上月下降'+str(abs(huanbi))+'个百分点'
    if huanbi == 0:
        judge3 = '较上月持平'
    return judge3 

# 计算环比差值
def calc_huanbizengzhang(df,period):
    this_month = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == period)]['数据值'])
    if period[5:6] == '01':
        last_period = str(period[0:4]-1) + '12'
    else:
        last_period = period[0:4] + str(int(period[5:7])-1).zfill(2)

    last_month = float(df[(df['地区编码'] == 110000) & (df['报告期编码'] == last_period)]['数据值'])
    huanbi = round(this_month-last_month,1)
    if huanbi > 0:
        judge5 = '较上月上升'+str(abs(huanbi))+'个百分点'
    if huanbi < 0:
        judge5 = '较上月下降'+str(abs(huanbi))+'个百分点'
    if huanbi == 0:
        judge5 = '较上月持平'
    return judge5


def get_paragraph1(df,period):
    text = ''
    value = str(float(df[(df['地区编码'] == 110000)&(df['报告期编码'] == period)]['数据值']))
    # 与全国对比结果
    judge1 = calc_country_change(df,period)
    # 同比结果
    judge2 = calc_tongbi(df,period)
    # 环比结果
    judge3 = calc_huanbi(df,period)
    
    text = period[0:4]+'年'+str(int(period[5:7]))+'月，北京市制造业采购经理指数（PMI）为'+\
           value+'%，'+judge1+'，'+judge2+'，'+judge3+'。'
    
    return text

def get_paragraph2(df,period):
    text = ''
    shengchan = df[df['产业名称'] == '制造业采购经理生产指数']
    xindingdan = df[df['产业名称'] == '制造业采购经理新订单指数']
    yuancailiao = df[df['产业名称'] == '制造业采购经理原材料库存指数']
    congyerenyuan = df[df['产业名称'] == '制造业采购经理从业人员指数']
    gongyingshang = df[df['产业名称'] == '制造业采购经理供应商配送时间指数']
    
    shengchan_value = str(round(float(shengchan[(shengchan['报告期编码']==period) & (shengchan['地区编码'] == 110000)]['数据值']),1))
    xindingdan_value = str(round(float(xindingdan[(xindingdan['报告期编码']==period) & (xindingdan['地区编码'] == 110000)]['数据值']),1))
    yuancailiao_value = str(round(float(yuancailiao[(yuancailiao['报告期编码']==period) & (yuancailiao['地区编码'] == 110000)]['数据值']),1))
    congyerenyuan_value = str(round(float(congyerenyuan[(congyerenyuan['报告期编码']==period) & (congyerenyuan['地区编码'] == 110000)]['数据值']),1))
    gongyingshang_value = str(round(float(gongyingshang[(gongyingshang['报告期编码']==period) & (gongyingshang['地区编码'] == 110000)]['数据值']),1))
    
    text = text+'从分项来看，生产指数为'+shengchan_value+'%，'+calc_huanbizengzhang(shengchan,period)+'；'+\
          '新订单指数为'+xindingdan_value+'%，'+calc_huanbizengzhang(xindingdan,period)+'；'+\
          '原材料库存指数为'+yuancailiao_value+'%，'+calc_huanbizengzhang(yuancailiao,period)+'；'+\
          '从业人员指数为'+congyerenyuan_value+'%，'+calc_huanbizengzhang(congyerenyuan,period)+'；'+\
          '供应商配送时间指数为'+gongyingshang_value+'%，'+calc_huanbizengzhang(gongyingshang,period)+'。'
          
    return text


######以下是段落设置

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING

from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

# 全文字体设置
def docxinitial(document):
    # 设置正文字体类型、大小
    document.styles["Normal"].font.name = u'Times New Roman'
    document.styles["Normal"].font.size = Pt(16) # 
    document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

# 大标题设置
def title0(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36) 
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(18)
    ti1.font.name = u'黑体'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 时间行设置
def time_note(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    ti2 = ti.add_run('\n')
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 一级标题设置
def title1(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 二级标题设置
def title2(document, txt):
    title = txt
    ti = document.add_paragraph()
    ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ti.paragraph_format.space_before = Pt(0)
    ti.paragraph_format.space_after = Pt(0)
    ti.paragraph_format.line_spacing = Pt(36)
    ti1 = ti.add_run(title)
    ti1.font.bold = True
    ti1.font.size = Pt(16)
    ti1.font.name = u'Times New Roman'
    r = ti1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    ti.paragraph_format.first_line_indent = ti.style.font.size * 2 # 缩进2字符
    ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

# 单一段落格式设置
def one_paragraph_txt(document, txt):
    txt1 = txt.split("，")[0]
    txt2 = txt.replace(txt1,'')
    p = document.add_paragraph()
    #p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(36)
    p.style.font.size = Pt(16)
    p1 = p.add_run(txt1)
    p1.font.size = Pt(16)
    p1.font.name = u'Times New Roman'
    r = p1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    p2 = p.add_run(txt2)
    p2.font.size = Pt(16)
    p.paragraph_format.first_line_indent = p.style.font.size * 2 # 缩进2字符
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距

    
# 多段格式设置
def much_paragraph_txt(document, txt):
    txt_list = txt.split("\n")
    for t in txt_list:
        one_paragraph_txt(document, t)

# 页码设置
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10.5)

def draw_line_chart(df):
    # 筛选出 2021 年及之后的数据
    df1=df.copy()
    df1['报告期编码'] = pd.to_datetime(df1['报告期编码'], format='%Y%m', errors='coerce')
    df1 = df1[df1['报告期编码'].dt.year >= 2021]
    
    #df1['date'] = df1['报告期编码'].apply(lambda x:int(x[0:4]))
    #df1 = df1[df1['date'] >= 2021]
    
    # 筛选出全国和北京的数据
    country1 = df1[(df1['地区编码'] == 0) | (df1['地区编码'] == '000000')][['报告期编码','数据值']]
    beijing1 = df1[df1['地区编码'] == 110000][['报告期编码','数据值']]
    
    # 创建图表
    plt.figure(figsize=(6,3.5))  # 设置图形大小
    plt.rcParams['font.sans-serif']=[u'SimHei']
    plt.rcParams['axes.unicode_minus']=False
    
    # 绘制全国数据，使用方形标记点
    plt.plot(country1['报告期编码'], country1['数据值'], label='全国', marker='s',markersize=4, linestyle='-',color='#4F81BD')
    # 绘制北京数据，使用方形标记点
    plt.plot(beijing1['报告期编码'], beijing1['数据值'], label='北京', marker='s',markersize=4, linestyle='-',color='#C0504D')
    # 设置图例、标题和坐标轴标签（使用中文）
    plt.legend(loc='lower center', ncol=2, frameon=False,bbox_to_anchor=(0.5, -0.2),)  # 添加图例，放在图的下方
# plt.title('全国和北京制造业 PMI指数月平均值折线图')  # 设置标题
# plt.xlabel('报告期编码')  # 设置 x 轴标签
# plt.ylabel('PMI指数_月平均值')  # 设置 y 轴标签
# plt.xticks(rotation=45)  # 设置 x 轴刻度角度
    
    # 获取 y 轴的最小和最大日期
    y_min = min(min(country1['数据值']), min(beijing1['数据值']))-2
    y_max = max(max(country1['数据值']), max(beijing1['数据值']))+2
    
    plt.ylim([y_min,y_max])
    
    # 设置 Y 轴标签格式为百分比
    def y_fmt(y, pos):
        return f'{y:.0f}%'

    plt.gca().yaxis.set_major_formatter(FuncFormatter(y_fmt))

    # 获取 x 轴的最小和最大日期
    x_min = min(min(country1['报告期编码']), min(beijing1['报告期编码']))
    x_max = max(max(country1['报告期编码']), max(beijing1['报告期编码']))

    # 获取 x 轴的刻度
    x_ticks = pd.date_range(start=x_min, end=x_max, freq='M').strftime('%Y-%m').tolist()
    x_ticks = [tick for tick in x_ticks if tick[-2:] in ['01', '04', '07', '10']]  # 保留 1、3、7、10 月份的刻度

    # 设置 X 轴的刻度和标签
    plt.xticks(x_ticks)
    plt.rcParams['xtick.labelsize'] = 8
    
    # 设置 X 轴的刻度为月份，间隔为3个月
    # plt.gca().xaxis.set_major_locator(MonthLocator(bymonthday=1, interval=3))
    date_format = DateFormatter("%Y%m")  # 设置日期格式为年月
    plt.gca().xaxis.set_major_formatter(date_format)
    

    plt.margins(0,0)
    plt.subplots_adjust(top=0.95,bottom=0.15,left=0.09,right=0.95)
    
    # 保存图表到 BytesIO 对象
    image_stream = BytesIO()
    plt.savefig(image_stream,format='png',bb_inches = 'tight',dpi=200)
    image_stream.seek(0)
   
    # 创建 Word 文档
    #doc = Document()
    
    # 将图表保存到 Word 文档中
    #doc.add_picture(image_stream)
    #doc.add_paragraph("图1：全国及北京制造业PMI")
    #doc.add_paragraph("数据来源：北京市统计局、国家统计局")

    # 保存 Word 文档
    #doc.save(r'.\报告输出\PMI月均值折线图.docx')
    
    return image_stream

df=pmi_other.copy()
period='202308'

#绘制雷达图 
def draw_radar_chart(df,period):
    matplotlib.rcParams['font.family'] = 'SimHei'
    matplotlib.rcParams['font.sans-serif'] = ['SimHei']
    
    shengchan = df[df['产业名称'] == '制造业采购经理生产指数']
    xindingdan = df[df['产业名称'] == '制造业采购经理新订单指数']
    yuancailiao = df[df['产业名称'] == '制造业采购经理原材料库存指数']
    congyerenyuan = df[df['产业名称'] == '制造业采购经理从业人员指数']
    gongyingshang = df[df['产业名称'] == '制造业采购经理供应商配送时间指数']
    
    shengchan_value = round(float(shengchan[(shengchan['报告期编码']==period) & (shengchan['地区编码'] == 110000)]['数据值']),1)
    xindingdan_value = round(float(xindingdan[(xindingdan['报告期编码']==period) & (xindingdan['地区编码'] == 110000)]['数据值']),1)
    yuancailiao_value = round(float(yuancailiao[(yuancailiao['报告期编码']==period) & (yuancailiao['地区编码'] == 110000)]['数据值']),1)
    congyerenyuan_value = round(float(congyerenyuan[(congyerenyuan['报告期编码']==period) & (congyerenyuan['地区编码'] == 110000)]['数据值']),1)
    gongyingshang_value = round(float(gongyingshang[(gongyingshang['报告期编码']==period) & (gongyingshang['地区编码'] == 110000)]['数据值']),1)
    
    labels = np.array(['生产指数','新订单指数','原材料库存指数','从业人员指数','供应商配送时间指数','生产指数'])
    values = np.array([shengchan_value,xindingdan_value,yuancailiao_value,congyerenyuan_value,gongyingshang_value])
    angles = np.linspace(0, 2*np.pi, 6)
    values = np.concatenate((values, [values[0]]))
    
    #plt.polar(angles,values)
    #plt.xticks(angles,labels)
    #plt.fill(angles,values)
    
    # 绘图
    fig=plt.figure(figsize=(6,3.5))
    # 这里一定要设置为极坐标格式
    ax = fig.add_subplot(111, polar=True)
    # 绘制折线图
    ax.plot(angles, values, '', linewidth=2)
    # 填充颜色
    ax.fill(angles, values, alpha=0.25)
    # 添加每个特征的标签
    ax.set_thetagrids(angles * 180/np.pi, labels,fontsize = 8)
    # 设置雷达图的范围
    minr = min(values)-5
    maxr = max(values)+5
    ax.set_rlim(minr,maxr)
    # 设置雷达图的0度起始位置
    ax.set_theta_zero_location('N')
    # 设置雷达图的坐标值显示角度，相对于起始角度的偏移量
    ax.set_rlabel_position(0)

    image_stream = BytesIO()
    plt.savefig(image_stream,format='png',bb_inches = 'tight',dpi=200)
    image_stream.seek(0)
    
    return image_stream


# 生成文件
def create_file(df,df2,period):
    period_str = period[0:4]+'年'+str(int(period[5:7]))+'月'
    
    document = Document()
    docxinitial(document)
    
    # 标题
    title0(document, '北京市制造业PMI分析月报')
    time_note(document, "（"+period_str+"）")
    
    #第一段落
    much_paragraph_txt(document, get_paragraph1(df,period))
    
    #第一张图
    image_stream=draw_line_chart(df)
    document.add_picture(image_stream)
    txt1='图1：全国及北京制造业PMI'
    p1=document.add_paragraph('')
    p1.alignment=WD_TABLE_ALIGNMENT.CENTER #居中
    run1 = p1.add_run(txt1)
    run1.font.size = Pt(11)
    txt2='数据来源：北京市统计局、国家统计局'
    p2=document.add_paragraph('')
    p2.alignment=WD_TABLE_ALIGNMENT.RIGHT #居右
    run2 = p2.add_run(txt2)
    run2.font.size = Pt(11)
    
    #第二段落
    much_paragraph_txt(document, get_paragraph2(df2,period))
    
    #第二张图
    image_stream2=draw_radar_chart(df2,period)
    document.add_picture(image_stream2)
    txt3='图2：制造业PMI各分项（%）'
    p1=document.add_paragraph('')
    p1.alignment=WD_TABLE_ALIGNMENT.CENTER #居中
    run1 = p1.add_run(txt3)
    run1.font.size = Pt(11)
    txt4='数据来源：北京市统计局、国家统计局'
    p2=document.add_paragraph('')
    p2.alignment=WD_TABLE_ALIGNMENT.RIGHT #居右
    run2 = p2.add_run(txt4)
    run2.font.size = Pt(11)
    
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] # 获取页眉的第一个段落 
    #paragraph.add_run('这是第一节的页眉') # 添加页面内容
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(paragraph.add_run())

    document.save(r'.\报告输出\北京市制造业PMI分析报告'+"（"+period_str+"）"+'.docx')

# 执行
create_file(pmi,pmi_other,'202309') 
create_file(pmi,pmi_other,'202308')  
create_file(pmi,pmi_other,'202307')
create_file(pmi,pmi_other,'202306')
create_file(pmi,pmi_other,'202305')
create_file(pmi,pmi_other,'202304')
create_file(pmi,pmi_other,'202303')
create_file(pmi,pmi_other,'202302')

