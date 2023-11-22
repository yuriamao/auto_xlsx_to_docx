import pandas as pd
from fuzzywuzzy import process
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
from tqdm import tqdm
from code.utils import *


plt.rcParams['font.sans-serif'] = ['Arial Unicode MS']  # 设置一个支持中文的字体，比如 Arial Unicode MS

def generate_monthly_chart(data_path, output_path, start_year, start_month):
    # 加载数据
    job_df = pd.read_excel(data_path, header=0, sheet_name='招聘')
    # 选择报告期编码为 '202307' 的数据
    selected_data = job_df.loc[job_df['指标名称'] == '地区分布岗位数量'].copy()
    selected_data['统计值'] = selected_data['统计值'].astype(int)
    # 按照集团系名称、报告期编码和招聘地区进行分组聚合
    aggregated_df = selected_data.groupby(['集团系名称', '报告期编码', '招聘地区']).agg({'统计值': 'sum'}).reset_index()
    # 将报告期编码转换为字符串格式
    aggregated_df['报告期编码'] = aggregated_df['报告期编码'].astype(str)
    # 添加年月份列，提取年份和月份前六个字符作为年月份
    aggregated_df['年月份'] = pd.to_datetime(aggregated_df['报告期编码'].str[:6], format='%Y%m')
    # 选择特定时间范围的数据
    filtered_data = aggregated_df[(aggregated_df['年月份'].dt.year > start_year) | 
                                  ((aggregated_df['年月份'].dt.year == start_year) & 
                                   (aggregated_df['年月份'].dt.month >= start_month))]

    # 按照年月份进行分组并对统计值进行求和
    grouped_data = filtered_data.groupby(['年月份', '招聘地区'])['统计值'].sum().reset_index()
    # 按年月份进行排序
    grouped_data = grouped_data.sort_values(by='年月份')
    # 创建一个新的 Figure 对象
    plt.figure(figsize=(10, 6))
    # 循环遍历不同的招聘地区
    for area in grouped_data['招聘地区'].unique():
        # 按照招聘地区筛选数据
        area_data = grouped_data[grouped_data['招聘地区'] == area]   
        # 绘制折线图，x轴为年月份，y轴为统计值
        plt.plot(area_data['年月份'], area_data['统计值'], label=area)

    # 添加图例
    plt.legend()
    # 添加标签和标题
    plt.xlabel('年月份')
    plt.ylabel('单位：个')
    # plt.title(f'信息软件业招聘岗位{start_year}年{start_month}月及之后不同招聘地区的月度统计值')
    # 显示图形
    # plt.xticks(rotation=45)  # 旋转x轴标签，使其更易读
    plt.tight_layout()  # 自动调整布局，避免标签重叠
    plt.savefig(output_path)  # 保存图形


def generate_report(date,img):
    # 创建一个新的 Word 文档
    doc = Document()
    # 添加标题并设置字体样式为宋体小四
    title = doc.add_heading('招聘统计信息', level=1)
    title_font = title.style.font
    title_font.name = '宋体'
    title_font.size = Pt(12)
    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'  # 设置表格样式
    # 添加表头并设置字体样式为宋体小四
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        cell.text = ''
    hdr_cells[3].text = '合计'
    hdr_cells[1].text = '本地'
    hdr_cells[2].text = '外地'

    # 遍历表格中所有的行和单元格，并设置字体样式为宋体小四
    for row in table.rows:
        for cell in row.cells:
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.name = '宋体'
            cell_font.size = Pt(12)
    
    # 重新加载数据并处理（根据实际数据路径和处理逻辑修改以下代码）

    job_df = pd.read_excel(r'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/指标数据20231120.xlsx', header=0, sheet_name='招聘')
    selected_data = job_df[(job_df['指标名称'] == '地区分布岗位数量')]
    selected_data.loc[:, '统计值'] = selected_data['统计值'].astype(int)
    # 按照集团系名称和招聘地区进行分组聚合统计值
    aggregated_df = selected_data.groupby(['集团系名称', '报告期编码', '招聘地区']).agg({'统计值': 'sum'}).reset_index()
    # 筛选出指定日期的数据
    selected_data = aggregated_df[aggregated_df['报告期编码'].astype(str).str.contains(date)]
    total_data = selected_data.groupby(['招聘地区']).agg({'统计值': 'sum'}).reset_index()
    local_val = total_data[total_data['招聘地区'] == '本地']['统计值'].values[0]
    foreign_val = total_data[total_data['招聘地区'] == '外地']['统计值'].values[0]
    total_val = local_val + foreign_val
    # 输出一段文字
    if not total_data.empty:
        output_text = f"{date[:4]}年{date[4:6]}月，北京市信息软件业招聘岗位数{total_val}个，其中本地{local_val}个，外地{foreign_val}个。"    
        time_year= f"（{date[:4]}年{date[4:6]}月）"
    else:
        output_text = "暂无数据。"
    # 按集团系名称进行聚合
    grouped_data = selected_data.groupby(['招聘地区', '集团系名称']).agg({'统计值': 'sum'}).reset_index()
    # 指定的 group_name 列表
    ordered_group_names = ['阿里系', '百度系', '抖音系', '京东系', '快手系', '美团系', '神州数码系', '小米系']
    # 获取 unique 集团系名称
    unique_group_names = selected_data['集团系名称'].unique()
    # 创建一个空字典，将模糊匹配的结果存储在其中
    matched_names = {}
    # 遍历每个 unique 集团系名称并找到最匹配的名称
    for name in unique_group_names:
        best_match = process.extractOne(name, ordered_group_names)[0]
        matched_names[name] = best_match
    # 根据模糊匹配的结果，对集团系名称进行重新排序
    selected_data = selected_data.copy()
    selected_data['集团系名称'] = selected_data['集团系名称'].map(matched_names)
    # 按集团系名称进行聚合
    grouped_data=selected_data.groupby(['招聘地区','集团系名称']).agg({'统计值': 'sum'}).reset_index()
    # 输出格式
    output = ""
    for group_name in ordered_group_names:
        group_data = grouped_data[grouped_data['集团系名称'] == group_name]
        if not group_data.empty:
            row = group_data.iloc[0]
            local_filtered_data = grouped_data[(grouped_data['集团系名称'] == row['集团系名称']) & (grouped_data['招聘地区'] == '本地')]
            foreign_filtered_data = grouped_data[(grouped_data['集团系名称'] == row['集团系名称']) & (grouped_data['招聘地区'] == '外地')]
            if not local_filtered_data.empty:
                local_value = int(local_filtered_data['统计值'].iloc[0])
            else:
                local_value = 0
            if not foreign_filtered_data.empty:
                foreign_value = int(foreign_filtered_data['统计值'].iloc[0])
            else:
                foreign_value = 0
            output += f"{group_name}招聘{local_value+foreign_value}人，其中本地{local_value}人，外地{foreign_value}人；"
            cells = table.add_row().cells
            cells[0].text = group_name # 集团系名称
            cells[3].text = str(local_value+foreign_value)  # 总招聘人数
            cells[1].text = str(local_value) # 本地招聘人数
            cells[2].text = str(foreign_value)  # 外地招聘人数
    prefix='分集团看，'
    output = output[:-1] + "。"
    par=prefix+output

    # 在 Word 文档中添加段落，包含输出文字
    # 插入图片
    image_path = img  # 图片路径

    doc.add_picture(image_path)  # 插入图片并设置宽度（可选）
    doc.add_paragraph(output_text)
    doc.add_paragraph(par)

    # 实例化 WordReportGenerator 类
    word_generator = WordReportGenerator()

    # 使用类方法创建文件
    word_generator.job_docx_file(title0='信息软件业招聘岗位数分析报告',
                             title1='',
                             par1=output_text,
                             pars=par,
                             file_pre='信息软件业招聘岗位数分析报告',
                             pic_dir='/Users/harvin/code/自动报告产品开发-产业链@20220830/data/chart_2022_nov.png',
                             file_dir='test',
                             year_month=time_year
                             )


    # 添加其他内容到文档...

    # 保存 Word 文档
    # doc.save(f'/Users/harvin/code/自动报告产品开发-产业链@20220830/data/output/招聘统计信息_{date}.docx')



def main():
    data_path = '/Users/harvin/code/自动报告产品开发-产业链@20220830/data/pmi_config.xlsx'
    output_image_path = '/Users/harvin/code/自动报告产品开发-产业链@20220830/data/chart_2022_nov.png'
    for i in tqdm(range(7,13)):
        generate_monthly_chart(data_path, output_image_path, start_year=2022, start_month=6)
        if i <10:
            i='0'+str(i)
        t='2022'+str(i)
        generate_report(t,output_image_path)

if __name__ == "__main__":
    # 执行 main 函数
    main()