from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime

class WordReportGenerator:
    def __init__(self):
        self.document = Document()
        self.docxinitial()

    def docxinitial(self):
        self.document.styles["Normal"].font.name = 'Times New Roman'
        self.document.styles["Normal"].font.size = Pt(16)

    def title0(self, txt):
        title = txt
        ti = self.document.add_paragraph()
        ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ti.paragraph_format.space_before = Pt(0)
        ti.paragraph_format.space_after = Pt(0)
        ti.paragraph_format.line_spacing = Pt(36) 
        ti1 = ti.add_run(title)
        ti1.font.size = Pt(18)
        ti1.font.name = '黑体'
        r = ti1._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def time_note(self, txt):
        title = txt
        ti = self.document.add_paragraph()
        ti.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ti.paragraph_format.space_before = Pt(0)
        ti.paragraph_format.space_after = Pt(0)
        ti.paragraph_format.line_spacing = Pt(36)
        ti1 = ti.add_run(title)
        ti1.font.size = Pt(16)
        ti1.font.name = 'Times New Roman'
        r = ti1._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        ti.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ... (other methods like title1, title2, one_paragraph_txt, much_paragraph_txt, create_element, create_attribute, add_page_number)

    def create_file(self, title0_name, date_str, title_txt, text_txt, output_file_path):
        time_str = datetime.datetime.strptime(date_str,'%Y-%m-%d')
        begin_date = (time_str-datetime.timedelta(days=time_str.weekday())).strftime("%Y年%m月%d日")
        end_date = (time_str+datetime.timedelta(days=6-time_str.weekday())).strftime("%Y年%m月%d日")
        
        self.title0(title0_name)
        self.time_note(f"（{begin_date}-{end_date}）")
        self.title1(title_txt)
        self.one_paragraph_txt(text_txt)
        
        # ... (other parts of your create_file method)

        self.document.save(output_file_path)
        print(f"File saved: {output_file_path}")

# Example usage:
report_generator = WordReportGenerator()
report_generator.create_file('Title', '2023-11-22', 'First Level Title', 'Sample text content', 'output.docx')
